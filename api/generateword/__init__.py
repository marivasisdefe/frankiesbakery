import logging
import azure.functions as func
from docx import Document
from io import BytesIO
 
def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Generando Word descargable desde salida del chat.')
 
    try:
        body = req.get_json()
        title = body.get("title", "Respuesta del Chatbot")
        author = body.get("author", "Usuario")
        content = body.get("content", "")
 
        # Crear documento Word
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Autor: {author}")
        doc.add_paragraph(content)
 
        # Guardar en memoria
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
 
        # Devolver archivo como descarga
        return func.HttpResponse(
            body=file_stream.getvalue(),
            status_code=200,
            headers={
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": 'attachment; filename="chat_output.docx"'
            }
        )
 
    except Exception as e:
        logging.error("Error en generateword: %s", e)
        return func.HttpResponse(
            f"Error al generar Word: {e}",
            status_code=500
        )
